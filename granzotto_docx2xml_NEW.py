#Questo script dà ancora problemi per le 
#forme {d::l} che nell'output vengono lette come grc
#anche se sono contenute in seg ita. 

!pip install python-docx
!pip install python-Levenshtein

import re
from google.colab import drive
drive.mount('/content/drive')

path_to_scripts = '/content/drive/Shareddrives/allographica/colab/scripts/'
path_to_docs = '/content/drive/Shareddrives/allographica/materiali_di_camilla_granzotto/'

from drive.Shareddrives.allographica.colab.scripts.mia.transcoder import Transcoder
from drive.Shareddrives.allographica.colab.scripts.mia.aligner import SimEval
from drive.Shareddrives.allographica.colab.scripts.mia.aligner import StringSimEval
from drive.Shareddrives.allographica.colab.scripts.mia.aligner import ObjectAligner

grc2lat = Transcoder(path_to_scripts+'grc2lat.txt')
grc2simpLat = Transcoder(path_to_scripts+'grc2simpLat.txt')
simpLat2regex = Transcoder(path_to_scripts+'simpLat2regex.txt')

from docx import Document

# Get the name of the uploaded file
in_file_name = path_to_docs + 'manuale_granzotto_unito.docx'
in_file = open(in_file_name, encoding = 'utf-8')

# Read the contents of the .docx file and divide by lines
doc = Document(in_file_name)

def create_para_list(doc):
    run_list = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.superscript:
                run_list.append(f'{{{run.text}}}')
            else:
                run_list.append(run.text)
        run_list.append('\n')
    text = ' '.join(run_list)
    return text.split('\n')

para_list = create_para_list(doc)
para_list.pop()

for i, para in enumerate(para_list):
    if i > 0 and i % 2 == 0:
        transl = grc2lat.transcode(para)

        simp_transl = grc2simpLat.transcode(para)

        # Verifica se 'simp_transl' è una lista
        if isinstance(simp_transl, list):
            # Unisci gli elementi della lista in una stringa
            simp_transl = ''.join(simp_transl)

        regex = simpLat2regex.transcode(simp_transl)

        print(para)
        print(transl)
        print(simp_transl)
        print(regex)

# Aggiunta di controllo per evitare IndexError
        if i + 1 < len(para_list):
            print(para_list[i+1])
        else:
            print("Nessun elemento successivo")

        print()

in_file.close()


#inizio teificazione
import re
from docx import Document
from lxml import etree
import xml.sax.saxutils as saxutils
from google.colab import files
import os

# Funzione aggiornata per fare escape di &, < e > con ‹ e ›
def escape_xml_chars(text):
    return text.replace('<', '‹').replace('>', '›').replace('&', '&amp;')

# Funzione per stabilire il testo per lingua
def determine_language(content):
    # Trattare come milestones forme come {δ::l}
    milestone_pattern = re.compile(r'\{[^\}]+\}')
    # Rimozione temporanea milestones per fare conteggio predominanza
    cleaned_content = milestone_pattern.sub('', content).strip()

    grc_pattern = re.compile(r'[Α-Ωα-ω]')
    ita_pattern = re.compile(r'[A-Za-z]')

    # Conteggio predominanza
    grc_count = len(grc_pattern.findall(cleaned_content))
    ita_count = len(ita_pattern.findall(cleaned_content))
    total_count = grc_count + ita_count

    if total_count == 0:
        return "empty"

    grc_percentage = (grc_count / total_count) * 100
    if grc_percentage > 20:  # Soglia del 20% per considerare greco il paragrafo
        return "grc"
    else:
        return "ita"

# Funzione per creare il file TEI
def create_tei_xml(paragraphs):
    TEI_NS = "http://www.tei-c.org/ns/1.0"
    XML_NS = "http://www.w3.org/XML/1998/namespace"
    NSMAP = {None: TEI_NS, 'xml': XML_NS}

    tei = etree.Element('TEI', nsmap=NSMAP)
    tei_header = etree.SubElement(tei, 'teiHeader')
    file_desc = etree.SubElement(tei_header, 'fileDesc')

    # Header
    title_stmt = etree.SubElement(file_desc, 'titleStmt')
    title = etree.SubElement(title_stmt, 'title')
    title.text = "Manuale per la costruzione delle navi"
    author = etree.SubElement(title_stmt, 'author')
    author.text = "Federico"
    publication_stmt = etree.SubElement(title_stmt, 'publicationStmt')
    pub = etree.SubElement(publication_stmt, 'p')
    pub.text = "Pubblicato da Camilla Granzotto, Federico Boschetti e Giulia Re"
    source_desc = etree.SubElement(file_desc, 'sourceDesc')
    desc = etree.SubElement(source_desc, 'p')
    desc.text = "Manoscritto allografico sulla costruzione delle navi"

    # Corpo
    text_element = etree.SubElement(tei, 'text')
    body_element = etree.SubElement(text_element, 'body')

    for para in paragraphs:
        body_element.append(para)

    return etree.tostring(tei, pretty_print=True, xml_declaration=True, encoding='UTF-8').decode('utf-8')

# Funzione per leggere e formattare il file .docx
def create_para_list_with_formatting(doc):
    para_list = []
    for para in doc.paragraphs:
        content = para.text.strip()
        if not content:
            continue

        # Riprende lang per stabilire la lingua del paragrafo
        lang = determine_language(content)

        # Creazione di <p> e segmentazione testo
        p_element = etree.Element("p")
        seg_element = etree.SubElement(p_element, "seg", type=lang)
        
        # Aggiunge i segmenti al paragrafo
        seg_element.text = escape_xml_chars(content)
        
        if seg_element.text.strip():
            para_list.append(p_element)
    
    return para_list

# Funzione per raffinare il linguaggio nell'XML
def refine_xml_language(xml_output):
    root = etree.fromstring(xml_output.encode('utf-8'))
    ns = {'tei': 'http://www.tei-c.org/ns/1.0'}

    # Itera su tutti gli elementi <seg>
    for seg in root.xpath("//tei:seg", namespaces=ns):
        text = seg.text or ""

        # Rimozione dei milestones per analisi
        milestone_pattern = re.compile(r'\{[^\}]+\}')
        cleaned_text = milestone_pattern.sub('', text).strip()
    
        # Conta caratteri greci e latini
        grc_count = len(re.findall(r'[Α-Ωα-ω]', cleaned_text))
        ita_count = len(re.findall(r'[A-Za-z]', cleaned_text))
        total_count = grc_count + ita_count

        # Calcola la percentuale di caratteri greci
        if total_count > 0:
            grc_percentage = (grc_count / total_count) * 100
            if grc_percentage <= 20:  # Soglia per correggere l'attributo
                seg.attrib['type'] = "ita"

    # Restituisci l'XML raffinato
    return etree.tostring(root, pretty_print=True, xml_declaration=True, encoding="UTF-8").decode("utf-8")

# Funzione principale
def main():
    # Percorso del file .docx
    input_docx = '/content/drive/Shareddrives/allographica/materiali_di_camilla_granzotto/manuale_granzotto_unito.docx'
    intermediate_xml = 'output_initial.xml'
    final_xml = 'output_final.xml'

    try:
        doc = Document(input_docx)
    except FileNotFoundError:
        print(f"Errore: il file {input_docx} non è stato trovato.")
        return None

    # Crea file XML iniziale
    para_list = create_para_list_with_formatting(doc)
    xml_output = create_tei_xml(para_list)
    
    # Salva l'XML iniziale
    with open(intermediate_xml, "w", encoding="utf-8") as f:
        f.write(xml_output)
    print(f"File iniziale salvato: {intermediate_xml}")

    # Raffina l'XML
    refined_xml = refine_xml_language(xml_output)

    # Salva il file XML finale
    with open(final_xml, "w", encoding="utf-8") as f:
        f.write(refined_xml)
    print(f"File finale salvato: {final_xml}")

    # Controlla se il file esiste prima di scaricarlo
    if os.path.exists(final_xml):
        print(f"File finale trovato: {final_xml}. Preparazione per il download...")
        files.download(final_xml)
    else:
        print(f"Errore: il file {final_xml} non è stato trovato.")
        return None

    return final_xml

# Esegui la funzione principale
main()
